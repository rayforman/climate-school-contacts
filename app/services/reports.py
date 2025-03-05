import os
import tempfile
from datetime import datetime
from flask import current_app
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import tempfile

from app.models import Event, Guest, EventAttendance

def generate_bio_sheet(event_id):
    """
    Generate a bio sheet for an event with the specified format.
    
    Args:
        event_id: ID of the event
        
    Returns:
        str: Path to the generated document
    """
    # Get the event and attendees (sorted by last name)
    event = Event.query.get_or_404(event_id)
    attendances = event.attendances.join(Guest).order_by(Guest.last_name, Guest.first_name).all()
    
    # Create a new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = f"Bio Sheet - {event.name}"
    doc.core_properties.author = "Columbia Climate School Contact Database"
    
    # Set default font to Georgia
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    
    # Add header with event info
    header = doc.add_paragraph()
    title_run = header.add_run(event.name)
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Format date and time
    event_date = event.date.strftime('%A, %B %d, %Y')
    event_time = event.date.strftime('%I:%M %p')
    if event.date.hour == 18 and event.date.minute == 0:  # Check if it's a standard 6-8pm event
        event_time = "6:00 - 8:00 PM"
    
    # Add date, time and location
    header.add_run(f"\n{event_date}\n{event_time}")
    if event.location:
        header.add_run(f"\n{event.location}")
    
    # Add spacer
    doc.add_paragraph()
    
    # Add each attendee
    for attendance in attendances:
        guest = attendance.guest
        
        # Create table for layout - 2 columns, one for photo and one for text
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        table.columns[0].width = Inches(1.2)  # Photo column
        table.columns[1].width = Inches(5.3)  # Text column
        
        # Get cells
        photo_cell = table.cell(0, 0)
        text_cell = table.cell(0, 1)
        
        # Add photo if available
        if guest.photo_filename:
            photo_path = os.path.join(current_app.config['UPLOAD_PATH'], guest.photo_filename)
            if os.path.exists(photo_path):
                try:
                    # Resize image for the document 
                    img = Image.open(photo_path)
                    
                    # Calculate height based on aspect ratio 
                    aspect_ratio = img.height / img.width
                    target_width = 1.1  # inches
                    target_height = target_width * aspect_ratio
                    
                    # Save to a temporary file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp:
                        img.save(temp.name, format='JPEG', quality=95)
                        photo_paragraph = photo_cell.paragraphs[0]
                        photo_run = photo_paragraph.add_run()
                        photo_run.add_picture(temp.name, width=Inches(target_width), height=Inches(target_height))
                        os.unlink(temp.name)  # Delete the temp file
                except Exception as e:
                    current_app.logger.error(f"Error adding photo: {str(e)}")
        
        # Add guest information to text cell
        
        # Name (Bold)
        name_paragraph = text_cell.paragraphs[0]
        name_run = name_paragraph.add_run(guest.full_name)
        name_run.bold = True
        name_run.font.size = Pt(12)
        
        # Athena ID (if available)
        if guest.athena_id:
            athena_paragraph = text_cell.add_paragraph()
            athena_paragraph.add_run(guest.athena_id)
        
        # Capacity and Prospect Manager (if available)
        if guest.donor_capacity or guest.prospect_manager:
            capacity_paragraph = text_cell.add_paragraph()
            
            if guest.donor_capacity and guest.prospect_manager:
                # Both exist, include the dash
                capacity_paragraph.add_run(f"{guest.donor_capacity} - {guest.prospect_manager}")
            elif guest.donor_capacity:
                # Only capacity exists
                capacity_paragraph.add_run(guest.donor_capacity)
            elif guest.prospect_manager:
                # Only prospect manager exists
                capacity_paragraph.add_run(guest.prospect_manager)
        
        # Bio (if available)
        if guest.bio:
            bio_paragraph = text_cell.add_paragraph()
            bio_paragraph.add_run(guest.bio)
        
        # Add empty paragraph as spacer between guests
        doc.add_paragraph()
    
    # Create a temporary file for the document
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        doc.save(temp_file.name)
        return temp_file.name